"""
[REP-02] DOCX reporter — genera reporte profesional en formato Word (.docx)
usando python-docx (sin dependencias de Node.js).
Soporta idiomas: 'en' (English) y 'es' (Español).
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

from reporter.base_reporter import BaseReporter
from schemas.controls_catalog import CONTROLS_BY_ID
from utils.logger import get_logger

logger = get_logger(__name__)

# ── Paleta de colores ─────────────────────────────────────────────────────────
C_PRIMARY   = RGBColor(0x1F, 0x38, 0x64)
C_SECONDARY = RGBColor(0x2E, 0x75, 0xB6)
C_ACCENT    = RGBColor(0x41, 0x71, 0x9C)
C_CRITICAL  = RGBColor(0x7F, 0x00, 0x00)
C_HIGH      = RGBColor(0xC0, 0x00, 0x00)
C_MEDIUM    = RGBColor(0xED, 0x7D, 0x31)
C_LOW       = RGBColor(0x70, 0xAD, 0x47)
C_PASS      = RGBColor(0x37, 0x56, 0x23)
C_FAIL      = RGBColor(0x7F, 0x00, 0x00)
C_SKIP      = RGBColor(0x59, 0x59, 0x59)
C_WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
C_LIGHTGRAY = RGBColor(0xF2, 0xF2, 0xF2)
C_DARKGRAY  = RGBColor(0x59, 0x59, 0x59)

SEV_COLORS = {
    "CRITICAL": C_CRITICAL,
    "HIGH":     C_HIGH,
    "MEDIUM":   C_MEDIUM,
    "LOW":      C_LOW,
    "INFO":     C_SECONDARY,
}

SEV_ORDER = ["CRITICAL", "HIGH", "MEDIUM", "LOW", "INFO"]

# ── Textos bilingüe ───────────────────────────────────────────────────────────
TEXTS = {
    "en": {
        "report_title":        "AWS SECURITY AUDIT REPORT",
        "account":             "Account",
        "region":              "Region",
        "generated":           "Generated",
        "auditor":             "Auditor",
        "score_label":         "Security Score",
        "total_controls":      "Total Controls",
        "passed":              "Passed",
        "failed":              "Failed",
        "skipped":             "Skipped",
        "framework_line":      "Framework: CIS AWS Foundations Benchmark v1.4  |  AWS Well-Architected Framework — Security Pillar",
        "confidential":        "CONFIDENTIAL — For authorized personnel only",
        "toc_title":           "Table of Contents",
        "s1_title":            "1. Executive Summary",
        "s1_intro":            "This report presents the results of an AWS security audit performed on account {account_id} (region: {region}) on {date}. The audit evaluated {total} security controls against the CIS AWS Foundations Benchmark v1.4 and the AWS Well-Architected Framework Security Pillar.",
        "s1_1":                "1.1 Overall Security Posture",
        "s1_1_body":           "The account achieved a global security score of {score}% — {label}. Out of {total} controls evaluated, {passed} passed, {failed} failed, and {skipped} were skipped.",
        "s1_2":                "1.2 Findings by Severity",
        "severity":            "Severity",
        "count":               "Count",
        "action":              "Action Required",
        "requires_attention":  "Requires attention",
        "no_findings":         "No findings",
        "s1_3":                "1.3 Critical Findings Requiring Immediate Attention",
        "critical_intro":      "The following {n} critical findings require immediate remediation:",
        "no_critical":         "No critical findings detected.",
        "s2_title":            "2. Methodology",
        "s2_1":                "2.1 Frameworks Applied",
        "fw1":                 "CIS AWS Foundations Benchmark v1.4 — Industry-standard security configuration guidelines for AWS accounts",
        "fw2":                 "AWS Well-Architected Framework — Security Pillar — AWS best practices for secure cloud architectures",
        "s2_2":                "2.2 Services Audited",
        "svc_line":            "{svc} — {total} controls evaluated ({passed} passed / {failed} failed)",
        "s2_3":                "2.3 Audit Scope",
        "s2_4":                "2.4 Severity Classification",
        "risk_level":          "Risk Level",
        "timeframe":           "Remediation Timeframe",
        "description":         "Description",
        "sev_rows": [
            ("CRITICAL", "Immediate",    "24–48 hours",       "Exploitation likely and impact severe. Treat as active incident."),
            ("HIGH",     "Urgent",       "1–2 weeks",         "High risk of exploitation. Prioritize in current sprint."),
            ("MEDIUM",   "Important",    "30 days",           "Moderate risk. Address in planned work."),
            ("LOW",      "Advisory",     "Next maintenance",  "Low risk. Address in next scheduled window."),
            ("INFO",     "Informational","As resources allow","Awareness item for posture improvement."),
        ],
        "s3_title":            "3. Results by Service",
        "s3_1":                "3.1 Service Score Summary",
        "score":               "Score",
        "total":               "Total",
        "svc_header":          "Service",
        "s3_svc":              "3.{n} {svc} — Score: {score}%",
        "svc_summary":         "{total} controls evaluated — {failed} failed, {passed} passed.",
        "failed_controls":     "Failed Controls ({n})",
        "control_id":          "Control ID",
        "finding":             "Finding",
        "remediation":         "Remediation",
        "s4_title":            "4. Prioritized Remediation Plan",
        "s4_intro":            "All failed controls ordered by severity. Address CRITICAL and HIGH items first — treat them as active security incidents.",
        "s4_1":                "4.1 Recommended Timeframes",
        "timeframes": [
            ("CRITICAL", "24–48 hours"),
            ("HIGH",     "1–2 weeks"),
            ("MEDIUM",   "30 days"),
            ("LOW",      "Next maintenance window"),
            ("INFO",     "As resources allow"),
        ],
        "s4_2":                "4.2 Remediation Table",
        "num":                 "#",
        "service":             "Service",
        "control":             "Control",
        "s5_title":            "5. Annexes",
        "s5_1":                "5.1 Passed Controls",
        "passed_intro":        "{n} controls passed.",
        "status":              "Status",
        "control_name":        "Control Name",
        "s5_2":                "5.2 Audit Metadata",
        "field":               "Field",
        "value":               "Value",
        "meta_rows":           [
            "Account ID", "Region", "Generated At", "Auditor",
            "Total Controls", "Global Score", "Framework 1", "Framework 2",
        ],
        "s5_3":                "5.3 Complete Controls Inventory",
        "catalog_intro":       "This section lists all {total} controls evaluated during the audit, showing their status (PASS / FAIL / SKIP), severity, and service.",
        "inventory_summary":   "Summary: {total} controls evaluated — {passed} PASS, {failed} FAIL, {skipped} SKIP.",
        "by_svc_header":       "Breakdown by Service",
        "s6_title":            "6. Reference Frameworks",
        "refs": [
            {
                "name": "CIS AWS Foundations Benchmark v1.4",
                "url":  "https://www.cisecurity.org/benchmark/amazon_web_services",
                "desc": "The CIS AWS Foundations Benchmark provides prescriptive guidance for establishing a secure baseline configuration for AWS. It defines controls across IAM, Storage, Logging, Monitoring, and Networking.",
            },
            {
                "name": "AWS Well-Architected Framework — Security Pillar",
                "url":  "https://docs.aws.amazon.com/wellarchitected/latest/security-pillar/welcome.html",
                "desc": "The Security Pillar of the AWS Well-Architected Framework describes how to take advantage of cloud technologies to protect data, systems, and assets. It covers identity management, detection, infrastructure protection, data protection, and incident response.",
            },
            {
                "name": "AWS Security Hub — CIS AWS Foundations Benchmark Standard",
                "url":  "https://docs.aws.amazon.com/securityhub/latest/userguide/cis-aws-foundations-benchmark.html",
                "desc": "AWS Security Hub's implementation of the CIS AWS Foundations Benchmark, providing automated compliance checks across your AWS environment.",
            },
            {
                "name": "NIST Cybersecurity Framework",
                "url":  "https://www.nist.gov/cyberframework",
                "desc": "The NIST CSF provides guidance on managing and reducing cybersecurity risk. Many CIS controls map directly to NIST CSF categories.",
            },
        ],
        "excellent": "Excellent", "good": "Good", "fair": "Fair",
        "poor": "Poor", "critical_lbl": "Critical",
    },
    "es": {
        "report_title":        "REPORTE DE AUDITORÍA DE SEGURIDAD AWS",
        "account":             "Cuenta",
        "region":              "Región",
        "generated":           "Generado",
        "auditor":             "Auditor",
        "score_label":         "Puntuación de Seguridad",
        "total_controls":      "Total de Controles",
        "passed":              "Aprobados",
        "failed":              "Fallidos",
        "skipped":             "Omitidos",
        "framework_line":      "Framework: CIS AWS Foundations Benchmark v1.4  |  AWS Well-Architected Framework — Pilar de Seguridad",
        "confidential":        "CONFIDENCIAL — Solo para personal autorizado",
        "toc_title":           "Tabla de Contenidos",
        "s1_title":            "1. Resumen Ejecutivo",
        "s1_intro":            "Este reporte presenta los resultados de una auditoría de seguridad AWS realizada en la cuenta {account_id} (región: {region}) el {date}. La auditoría evaluó {total} controles de seguridad contra el CIS AWS Foundations Benchmark v1.4 y el AWS Well-Architected Framework — Pilar de Seguridad.",
        "s1_1":                "1.1 Postura de Seguridad General",
        "s1_1_body":           "La cuenta obtuvo una puntuación global de seguridad de {score}% — {label}. De los {total} controles evaluados, {passed} aprobaron, {failed} fallaron y {skipped} fueron omitidos.",
        "s1_2":                "1.2 Hallazgos por Severidad",
        "severity":            "Severidad",
        "count":               "Cantidad",
        "action":              "Acción Requerida",
        "requires_attention":  "Requiere atención",
        "no_findings":         "Sin hallazgos",
        "s1_3":                "1.3 Hallazgos Críticos que Requieren Atención Inmediata",
        "critical_intro":      "Los siguientes {n} hallazgos críticos requieren remediación inmediata:",
        "no_critical":         "No se detectaron hallazgos críticos.",
        "s2_title":            "2. Metodología",
        "s2_1":                "2.1 Frameworks Aplicados",
        "fw1":                 "CIS AWS Foundations Benchmark v1.4 — Guía prescriptiva estándar para configuración segura de cuentas AWS",
        "fw2":                 "AWS Well-Architected Framework — Pilar de Seguridad — Mejores prácticas de AWS para arquitecturas cloud seguras",
        "s2_2":                "2.2 Servicios Auditados",
        "svc_line":            "{svc} — {total} controles evaluados ({passed} aprobados / {failed} fallidos)",
        "s2_3":                "2.3 Alcance de la Auditoría",
        "s2_4":                "2.4 Clasificación de Severidades",
        "risk_level":          "Nivel de Riesgo",
        "timeframe":           "Plazo de Remediación",
        "description":         "Descripción",
        "sev_rows": [
            ("CRITICAL", "Inmediato",       "24–48 horas",            "Explotación probable e impacto severo. Tratar como incidente activo."),
            ("HIGH",     "Urgente",         "1–2 semanas",            "Alto riesgo de explotación. Priorizar en el sprint actual."),
            ("MEDIUM",   "Importante",      "30 días",                "Riesgo moderado. Abordar en el trabajo planificado."),
            ("LOW",      "Informativo",     "Próximo mantenimiento",  "Riesgo bajo. Abordar en la próxima ventana de mantenimiento."),
            ("INFO",     "Informacional",   "Según disponibilidad",   "Elemento de concienciación para mejorar la postura de seguridad."),
        ],
        "s3_title":            "3. Resultados por Servicio",
        "s3_1":                "3.1 Resumen de Puntuación por Servicio",
        "score":               "Puntuación",
        "total":               "Total",
        "svc_header":          "Servicio",
        "s3_svc":              "3.{n} {svc} — Puntuación: {score}%",
        "svc_summary":         "{total} controles evaluados — {failed} fallidos, {passed} aprobados.",
        "failed_controls":     "Controles Fallidos ({n})",
        "control_id":          "ID de Control",
        "finding":             "Hallazgo",
        "remediation":         "Remediación",
        "s4_title":            "4. Plan de Remediación Priorizado",
        "s4_intro":            "Todos los controles fallidos ordenados por severidad. Abordar los elementos CRITICAL y HIGH primero — tratarlos como incidentes de seguridad activos.",
        "s4_1":                "4.1 Plazos Recomendados",
        "timeframes": [
            ("CRITICAL", "24–48 horas"),
            ("HIGH",     "1–2 semanas"),
            ("MEDIUM",   "30 días"),
            ("LOW",      "Próxima ventana de mantenimiento"),
            ("INFO",     "Según disponibilidad de recursos"),
        ],
        "s4_2":                "4.2 Tabla de Remediación",
        "num":                 "#",
        "service":             "Servicio",
        "control":             "Control",
        "s5_title":            "5. Anexos",
        "s5_1":                "5.1 Controles Aprobados",
        "passed_intro":        "{n} controles aprobados.",
        "status":              "Estado",
        "control_name":        "Nombre del Control",
        "s5_2":                "5.2 Metadatos de la Auditoría",
        "field":               "Campo",
        "value":               "Valor",
        "meta_rows":           [
            "ID de Cuenta", "Región", "Generado En", "Auditor",
            "Total de Controles", "Puntuación Global", "Framework 1", "Framework 2",
        ],
        "s5_3":                "5.3 Inventario Completo de Controles",
        "catalog_intro":       "Esta sección lista los {total} controles evaluados durante la auditoría, mostrando su estado (PASS / FAIL / SKIP), severidad y servicio.",
        "inventory_summary":   "Resumen: {total} controles evaluados — {passed} PASS, {failed} FAIL, {skipped} SKIP.",
        "by_svc_header":       "Desglose por Servicio",
        "s6_title":            "6. Frameworks de Referencia",
        "refs": [
            {
                "name": "CIS AWS Foundations Benchmark v1.4",
                "url":  "https://www.cisecurity.org/benchmark/amazon_web_services",
                "desc": "El CIS AWS Foundations Benchmark proporciona guía prescriptiva para establecer una configuración base segura en AWS. Define controles en IAM, Almacenamiento, Logging, Monitoreo y Networking.",
            },
            {
                "name": "AWS Well-Architected Framework — Pilar de Seguridad",
                "url":  "https://docs.aws.amazon.com/wellarchitected/latest/security-pillar/welcome.html",
                "desc": "El Pilar de Seguridad del AWS Well-Architected Framework describe cómo aprovechar las tecnologías cloud para proteger datos, sistemas y activos. Cubre gestión de identidades, detección, protección de infraestructura, protección de datos y respuesta a incidentes.",
            },
            {
                "name": "AWS Security Hub — Estándar CIS AWS Foundations Benchmark",
                "url":  "https://docs.aws.amazon.com/securityhub/latest/userguide/cis-aws-foundations-benchmark.html",
                "desc": "Implementación de AWS Security Hub del CIS AWS Foundations Benchmark, con verificaciones automatizadas de cumplimiento en todo el entorno AWS.",
            },
            {
                "name": "NIST Cybersecurity Framework",
                "url":  "https://www.nist.gov/cyberframework",
                "desc": "El NIST CSF provee guía para gestionar y reducir el riesgo de ciberseguridad. Muchos controles CIS se mapean directamente a categorías del NIST CSF.",
            },
        ],
        "excellent": "Excelente", "good": "Bueno", "fair": "Regular",
        "poor": "Deficiente", "critical_lbl": "Crítico",
    },
}


def _score_color(score: float) -> RGBColor:
    if score >= 90: return C_LOW
    if score >= 75: return RGBColor(0x53, 0x81, 0x35)
    if score >= 50: return C_MEDIUM
    if score >= 25: return C_HIGH
    return C_CRITICAL


def _score_label(score: float, t: dict) -> str:
    if score >= 90: return t["excellent"]
    if score >= 75: return t["good"]
    if score >= 50: return t["fair"]
    if score >= 25: return t["poor"]
    return t["critical_lbl"]


# ── Helpers de XML para shading y bordes ──────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    """Aplica color de fondo a una celda."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def _set_cell_border(cell, color: str = "CCCCCC"):
    """Aplica borde fino a todos los lados de una celda."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _set_row_height(row, height_cm: float):
    tr   = row._tr
    trPr = tr.get_or_add_trPr()
    trH  = OxmlElement('w:trHeight')
    trH.set(qn('w:val'), str(int(height_cm * 567)))
    trPr.append(trH)


def _para_border_bottom(para, hex_color: str, size: int = 12):
    """Agrega borde inferior a un párrafo."""
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    str(size))
    bot.set(qn('w:space'), '4')
    bot.set(qn('w:color'), hex_color)
    pBdr.append(bot)
    pPr.append(pBdr)


class DocxReporter(BaseReporter):

    def __init__(self, report, output_dir: str, timestamp: str, lang: str = "en"):
        super().__init__(report, output_dir, timestamp)
        self.lang = lang if lang in TEXTS else "en"
        self.t    = TEXTS[self.lang]

    def _output_path(self, extension: str):
        """Incluye el idioma en el nombre del archivo."""
        return self.output_dir / f"aws_audit_{self.report.account_id}_{self.timestamp}_{self.lang}.{extension}"

    def generate(self) -> str:
        logger.info(f"Generating DOCX report [{self.lang.upper()}]...")
        path = self._output_path("docx")
        path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()
        self._setup_document(doc)
        self._build_cover(doc)
        self._build_toc(doc)
        self._build_executive_summary(doc)
        self._build_methodology(doc)
        self._build_service_scores(doc)
        self._build_service_findings(doc)
        self._build_remediation_plan(doc)
        self._build_annex(doc)
        self._build_references(doc)

        doc.save(str(path))
        logger.info(f"DOCX report saved: {path}")
        return str(path)

    # ── Setup ──────────────────────────────────────────────────────────────────

    def _setup_document(self, doc: Document):
        """Configura márgenes y estilos base del documento."""
        section = doc.sections[0]
        section.page_width   = Inches(8.5)
        section.page_height  = Inches(11)
        section.left_margin  = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin   = Inches(1)
        section.bottom_margin = Inches(1)

        # Estilo de fuente por defecto
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

    # ── Helpers de contenido ───────────────────────────────────────────────────

    def _h1(self, doc: Document, text: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(6)
        run = p.add_run(text)
        run.font.name  = 'Calibri'
        run.font.size  = Pt(18)
        run.font.bold  = True
        run.font.color.rgb = C_PRIMARY
        _para_border_bottom(p, "2E75B6", 12)
        return p

    def _h2(self, doc: Document, text: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run(text)
        run.font.name  = 'Calibri'
        run.font.size  = Pt(14)
        run.font.bold  = True
        run.font.color.rgb = C_SECONDARY
        return p

    def _h3(self, doc: Document, text: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(text)
        run.font.name  = 'Calibri'
        run.font.size  = Pt(12)
        run.font.bold  = True
        run.font.color.rgb = C_ACCENT
        return p

    def _body(self, doc: Document, text: str, bold=False, color=None, italic=False, size=11):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(text)
        run.font.name   = 'Calibri'
        run.font.size   = Pt(size)
        run.font.bold   = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = color
        return p

    def _bullet(self, doc: Document, text: str, color=None):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        if color:
            run.font.color.rgb = color
        return p

    def _page_break(self, doc: Document):
        doc.add_page_break()

    def _make_table(self, doc: Document, headers: list, col_widths: list) -> object:
        """Crea tabla con cabecera formateada."""
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'

        # Fila de cabecera
        hdr_row = table.rows[0]
        for i, (hdr, width) in enumerate(zip(headers, col_widths)):
            cell = hdr_row.cells[i]
            cell.width = Inches(width)
            _set_cell_bg(cell, "1F3864")
            _set_cell_border(cell, "1F3864")
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            run = p.add_run(hdr)
            run.font.name  = 'Calibri'
            run.font.size  = Pt(10)
            run.font.bold  = True
            run.font.color.rgb = C_WHITE

        return table

    def _add_data_row(self, table, values: list, col_widths: list,
                      row_idx: int = 0, colors: list = None):
        """Agrega fila de datos a la tabla."""
        row  = table.add_row()
        bg   = "F2F2F2" if row_idx % 2 == 0 else "FFFFFF"

        for i, (val, width) in enumerate(zip(values, col_widths)):
            cell = row.cells[i]
            cell.width = Inches(width)
            _set_cell_bg(cell, bg)
            _set_cell_border(cell, "CCCCCC")
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            run = p.add_run(str(val or ""))
            run.font.name = 'Calibri'
            run.font.size = Pt(9)
            if colors and colors[i]:
                run.font.color.rgb = colors[i]
                run.font.bold = True

        return row

    # ── Secciones ──────────────────────────────────────────────────────────────
    # ── Secciones ──────────────────────────────────────────────────────────────

    def _build_cover(self, doc: Document):
        r = self.report
        t = self.t

        for _ in range(4):
            doc.add_paragraph()

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(t["report_title"])
        run.font.name = 'Calibri'; run.font.size = Pt(28)
        run.font.bold = True; run.font.color.rgb = C_PRIMARY
        _para_border_bottom(p, "2E75B6", 16)

        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{t['account']}: {r.account_id}  |  {t['region']}: {r.region}")
        run.font.name = 'Calibri'; run.font.size = Pt(13); run.font.color.rgb = C_DARKGRAY

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{t['generated']}: {r.generated_at.strftime('%Y-%m-%d %H:%M UTC')}  |  {t['auditor']}: {r.auditor or 'N/A'}")
        run.font.name = 'Calibri'; run.font.size = Pt(11)
        run.font.italic = True; run.font.color.rgb = C_DARKGRAY

        doc.add_paragraph()
        score = r.global_score
        score_clr = _score_color(score)
        score_lbl = _score_label(score, t)

        tbl = doc.add_table(rows=2, cols=1)
        tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c1 = tbl.rows[0].cells[0]
        c1.width = Inches(3); _set_cell_bg(c1, "EEF3F8"); _set_cell_border(c1, "2E75B6")
        p = c1.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{score}%")
        run.font.name = 'Calibri'; run.font.size = Pt(48); run.font.bold = True; run.font.color.rgb = score_clr

        c2 = tbl.rows[1].cells[0]
        c2.width = Inches(3); _set_cell_bg(c2, "EEF3F8"); _set_cell_border(c2, "2E75B6")
        p = c2.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{t['score_label']} — {score_lbl}")
        run.font.name = 'Calibri'; run.font.size = Pt(13); run.font.color.rgb = C_DARKGRAY

        doc.add_paragraph()
        stats_tbl = doc.add_table(rows=2, cols=4)
        stats_tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdrs = [t["total_controls"], f"✓ {t['passed']}", f"✗ {t['failed']}", f"— {t['skipped']}"]
        vals = [r.total_controls, r.passed_controls, r.failed_controls, r.skipped_controls]
        val_colors = [C_PRIMARY, C_PASS, C_FAIL, C_SKIP]
        hdr_bgs    = ["41719C", "375623", "7F0000", "595959"]

        for i, (hdr, bg) in enumerate(zip(hdrs, hdr_bgs)):
            cell = stats_tbl.rows[0].cells[i]
            _set_cell_bg(cell, bg); _set_cell_border(cell, bg)
            p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(hdr)
            run.font.name = 'Calibri'; run.font.size = Pt(10); run.font.bold = True; run.font.color.rgb = C_WHITE

        for i, (val, clr) in enumerate(zip(vals, val_colors)):
            cell = stats_tbl.rows[1].cells[i]
            _set_cell_bg(cell, "F2F2F2"); _set_cell_border(cell, "CCCCCC")
            p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.name = 'Calibri'; run.font.size = Pt(20); run.font.bold = True; run.font.color.rgb = clr

        doc.add_paragraph()
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(t["framework_line"])
        run.font.name = 'Calibri'; run.font.size = Pt(9); run.font.italic = True; run.font.color.rgb = C_DARKGRAY

        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(t["confidential"])
        run.font.name = 'Calibri'; run.font.size = Pt(10); run.font.bold = True; run.font.color.rgb = C_CRITICAL

        self._page_break(doc)

    def _build_toc(self, doc: Document):
        t = self.t
        self._h1(doc, t["toc_title"])
        toc_entries = [
            (1, t["s1_title"]), (2, t["s1_1"]), (2, t["s1_2"]), (2, t["s1_3"]),
            (1, t["s2_title"]), (2, t["s2_1"]), (2, t["s2_2"]), (2, t["s2_3"]), (2, t["s2_4"]),
            (1, t["s3_title"]), (2, t["s3_1"]),
            (1, t["s4_title"]),
            (1, t["s5_title"]), (2, t["s5_1"]), (2, t["s5_2"]), (2, t["s5_3"]),
            (1, t["s6_title"]),
        ]
        for level, text in toc_entries:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3 if level == 1 else 1)
            p.paragraph_format.space_after  = Pt(1)
            if level == 2:
                p.paragraph_format.left_indent = Inches(0.3)
            run = p.add_run(text)
            run.font.name = 'Calibri'; run.font.size = Pt(11 if level == 1 else 10)
            run.font.bold = level == 1
            run.font.color.rgb = C_PRIMARY if level == 1 else C_DARKGRAY
        self._page_break(doc)

    def _build_executive_summary(self, doc: Document):
        r = self.report; t = self.t
        self._h1(doc, t["s1_title"])
        self._body(doc, t["s1_intro"].format(
            account_id=r.account_id, region=r.region,
            date=r.generated_at.strftime('%Y-%m-%d'),
            total=r.total_controls,
        ))
        self._h2(doc, t["s1_1"])
        self._body(doc, t["s1_1_body"].format(
            score=r.global_score, label=_score_label(r.global_score, t),
            total=r.total_controls, passed=r.passed_controls,
            failed=r.failed_controls, skipped=r.skipped_controls,
        ))
        self._h2(doc, t["s1_2"])
        tbl = self._make_table(doc, [t["severity"], t["count"], t["action"]], [1.5, 1.0, 4.0])
        for i, sev in enumerate(SEV_ORDER):
            count = r.by_severity.get(sev, 0)
            self._add_data_row(tbl, [sev, str(count),
                t["requires_attention"] if count > 0 else t["no_findings"]],
                [1.5, 1.0, 4.0], row_idx=i,
                colors=[SEV_COLORS.get(sev), None, SEV_COLORS.get(sev) if count > 0 else C_PASS],
            )
        self._h2(doc, t["s1_3"])
        critical = [f for f in r.findings if f.severity.value == "CRITICAL" and f.status.value == "FAIL"]
        if critical:
            self._body(doc, t["critical_intro"].format(n=len(critical)), bold=True)
            for f in critical[:10]:
                self._bullet(doc, f"[{f.service.upper()}] {f.control_id}: {f.message[:150]}", color=C_CRITICAL)
        else:
            self._body(doc, t["no_critical"], color=C_PASS)
        self._page_break(doc)

    def _build_methodology(self, doc: Document):
        r = self.report; t = self.t
        self._h1(doc, t["s2_title"])
        self._h2(doc, t["s2_1"])
        self._bullet(doc, t["fw1"]); self._bullet(doc, t["fw2"])
        self._h2(doc, t["s2_2"])
        for s in r.by_service:
            self._bullet(doc, t["svc_line"].format(
                svc=s.service.upper(), total=s.total, passed=s.passed, failed=s.failed))
        self._h2(doc, t["s2_3"])
        self._body(doc, f"{t['account']} ID:  {r.account_id}")
        self._body(doc, f"{t['region']}:      {r.region}")
        self._body(doc, f"{t['generated']}:   {r.generated_at.strftime('%Y-%m-%d %H:%M UTC')}")
        self._body(doc, f"{t['auditor']}:     {r.auditor or 'N/A'}")
        self._h2(doc, t["s2_4"])
        tbl = self._make_table(doc,
            [t["severity"], t["risk_level"], t["timeframe"], t["description"]],
            [1.0, 1.2, 1.8, 2.5])
        for i, (sev, risk, time, desc) in enumerate(t["sev_rows"]):
            self._add_data_row(tbl, [sev, risk, time, desc], [1.0, 1.2, 1.8, 2.5],
                               row_idx=i, colors=[SEV_COLORS.get(sev), None, None, None])
        self._page_break(doc)

    def _build_service_scores(self, doc: Document):
        r = self.report; t = self.t
        self._h1(doc, t["s3_title"])
        self._h2(doc, t["s3_1"])
        tbl = self._make_table(doc,
            [t["svc_header"], t["score"], t["total"], t["passed"], t["failed"], t["skipped"]],
            [1.2, 0.8, 0.8, 0.9, 0.9, 0.9])
        for i, s in enumerate(r.by_service):
            self._add_data_row(tbl,
                [s.service.upper(), f"{s.score}%", s.total, s.passed, s.failed, s.skipped],
                [1.2, 0.8, 0.8, 0.9, 0.9, 0.9], row_idx=i,
                colors=[C_PRIMARY, _score_color(s.score), None, C_PASS, C_FAIL if s.failed > 0 else None, None])

    def _build_service_findings(self, doc: Document):
        r = self.report; t = self.t
        by_service = {}
        for f in r.findings:
            by_service.setdefault(f.service, []).append(f)
        svc_order = ["iam", "s3", "ec2", "rds", "guardduty", "cloudwatch", "vpc"]
        services  = [s for s in svc_order if s in by_service] + \
                    [s for s in by_service if s not in svc_order]
        for idx, svc in enumerate(services):
            findings  = by_service[svc]
            svc_score = next((s.score for s in r.by_service if s.service == svc), 0)
            failed    = sorted([f for f in findings if f.status.value == "FAIL"],
                               key=lambda x: SEV_ORDER.index(x.severity.value))
            self._h2(doc, t["s3_svc"].format(n=idx + 2, svc=svc.upper(), score=svc_score))
            self._body(doc, t["svc_summary"].format(
                total=len(findings), failed=len(failed), passed=len(findings)-len(failed)),
                italic=True, color=C_DARKGRAY)
            if failed:
                self._h3(doc, t["failed_controls"].format(n=len(failed)))
                tbl = self._make_table(doc,
                    [t["control_id"], t["severity"], t["finding"], t["remediation"]],
                    [1.2, 0.9, 2.9, 1.5])
                for i, f in enumerate(failed):
                    # Usar nombre del control en el idioma correcto si está en el catálogo
                    catalog_entry = CONTROLS_BY_ID.get(f.control_id, {})
                    name_key = f"name_{self.lang}"
                    ctrl_name = catalog_entry.get(name_key, f.control_name)
                    self._add_data_row(tbl,
                        [f.control_id, f.severity.value, f.message[:200], f.remediation[:150]],
                        [1.2, 0.9, 2.9, 1.5], row_idx=i,
                        colors=[None, SEV_COLORS.get(f.severity.value), None, None])
                doc.add_paragraph()
        self._page_break(doc)

    def _build_remediation_plan(self, doc: Document):
        r = self.report; t = self.t
        self._h1(doc, t["s4_title"])
        self._body(doc, t["s4_intro"])
        self._h2(doc, t["s4_1"])
        for sev, time in t["timeframes"]:
            self._bullet(doc, f"{sev}: {time}", color=SEV_COLORS.get(sev))
        self._h2(doc, t["s4_2"])
        seen = set(); all_failed = []
        for f in sorted(r.findings, key=lambda x: SEV_ORDER.index(x.severity.value)):
            if f.status.value == "FAIL":
                key = f"{f.control_id}-{f.service}"
                if key not in seen:
                    seen.add(key); all_failed.append(f)
        tbl = self._make_table(doc,
            [t["num"], t["severity"], t["service"], t["control"], t["action"]],
            [0.3, 0.9, 0.7, 1.2, 3.4])
        for i, f in enumerate(all_failed):
            self._add_data_row(tbl,
                [str(i+1), f.severity.value, f.service.upper(), f.control_id, f.remediation[:250]],
                [0.3, 0.9, 0.7, 1.2, 3.4], row_idx=i,
                colors=[None, SEV_COLORS.get(f.severity.value), None, None, None])
        self._page_break(doc)

    def _build_annex(self, doc: Document):
        r = self.report; t = self.t
        self._h1(doc, t["s5_title"])

        # 5.1 Passed controls — TODOS, sin límite
        self._h2(doc, t["s5_1"])
        passed = [f for f in r.findings if f.status.value == "PASS"]
        self._body(doc, t["passed_intro"].format(n=len(passed)), italic=True, color=C_DARKGRAY)
        tbl = self._make_table(doc,
            [t["control_id"], t["service"], t["status"], t["control_name"]],
            [1.5, 0.8, 0.7, 3.5])
        for i, f in enumerate(passed):
            catalog_entry = CONTROLS_BY_ID.get(f.control_id, {})
            name_key = f"name_{self.lang}"
            ctrl_name = catalog_entry.get(name_key, f.control_name)
            self._add_data_row(tbl,
                [f.control_id, f.service.upper(), "PASS", ctrl_name],
                [1.5, 0.8, 0.7, 3.5], row_idx=i,
                colors=[None, None, C_PASS, None])

        self._page_break(doc)

        # 5.2 Metadata
        self._h2(doc, t["s5_2"])
        meta_keys   = t["meta_rows"]
        meta_values = [
            r.account_id,
            r.region,
            r.generated_at.strftime('%Y-%m-%d %H:%M UTC'),
            r.auditor or "N/A",
            str(r.total_controls),
            f"{r.global_score}%",
            "CIS AWS Foundations Benchmark v1.4",
            "AWS Well-Architected Framework — Security Pillar",
        ]
        tbl = self._make_table(doc, [t["field"], t["value"]], [2.0, 4.5])
        for i, (k, v) in enumerate(zip(meta_keys, meta_values)):
            self._add_data_row(tbl, [k, v], [2.0, 4.5], row_idx=i, colors=[C_PRIMARY, None])

        self._page_break(doc)

        # 5.3 Inventario Completo de Controles — TODOS los findings del JSON
        self._h2(doc, t["s5_3"])

        total   = len(r.findings)
        n_pass  = sum(1 for f in r.findings if f.status.value == "PASS")
        n_fail  = sum(1 for f in r.findings if f.status.value == "FAIL")
        n_skip  = sum(1 for f in r.findings if f.status.value == "SKIP")

        self._body(doc, t["catalog_intro"].format(total=total), italic=True, color=C_DARKGRAY)
        doc.add_paragraph()

        # Resumen numérico
        p = doc.add_paragraph()
        run = p.add_run(t["inventory_summary"].format(
            total=total, passed=n_pass, failed=n_fail, skipped=n_skip))
        run.font.name = 'Calibri'; run.font.size = Pt(12); run.font.bold = True
        run.font.color.rgb = C_PRIMARY

        doc.add_paragraph()

        # Desglose por servicio
        self._h3(doc, t["by_svc_header"])
        svc_stats = {}
        for f in r.findings:
            svc = f.service.upper()
            if svc not in svc_stats:
                svc_stats[svc] = {"total": 0, "pass": 0, "fail": 0, "skip": 0}
            svc_stats[svc]["total"] += 1
            if f.status.value == "PASS":
                svc_stats[svc]["pass"] += 1
            elif f.status.value == "FAIL":
                svc_stats[svc]["fail"] += 1
            else:
                svc_stats[svc]["skip"] += 1

        svc_order = ["IAM", "S3", "EC2", "RDS", "GUARDDUTY", "CLOUDWATCH", "VPC"]
        sorted_svcs = [s for s in svc_order if s in svc_stats] + \
                      [s for s in svc_stats if s not in svc_order]

        tbl = self._make_table(doc,
            [t["service"], t["total"], t["passed"], t["failed"], t["skipped"]],
            [1.5, 0.8, 0.8, 0.8, 0.8])
        for i, svc in enumerate(sorted_svcs):
            st = svc_stats[svc]
            self._add_data_row(tbl,
                [svc, st["total"], st["pass"], st["fail"], st["skip"]],
                [1.5, 0.8, 0.8, 0.8, 0.8], row_idx=i,
                colors=[C_PRIMARY, None, C_PASS if st["pass"] > 0 else None,
                        C_FAIL if st["fail"] > 0 else None, None])

        self._page_break(doc)

        # Tabla completa de todos los controles, ordenados por servicio y luego por status
        STATUS_COLORS = {"PASS": C_PASS, "FAIL": C_FAIL, "SKIP": C_SKIP}

        all_sorted = sorted(r.findings, key=lambda f: (
            svc_order.index(f.service.upper()) if f.service.upper() in svc_order else 99,
            0 if f.status.value == "FAIL" else (1 if f.status.value == "SKIP" else 2),
            SEV_ORDER.index(f.severity.value) if f.severity.value in SEV_ORDER else 99,
        ))

        tbl = self._make_table(doc,
            [t["control_id"], t["service"], t["severity"], t["status"], t["control_name"]],
            [1.3, 0.7, 0.7, 0.6, 3.2])

        for i, f in enumerate(all_sorted):
            catalog_entry = CONTROLS_BY_ID.get(f.control_id, {})
            name_key = f"name_{self.lang}"
            ctrl_name = catalog_entry.get(name_key, f.control_name)
            status_clr = STATUS_COLORS.get(f.status.value, C_DARKGRAY)
            self._add_data_row(tbl,
                [f.control_id, f.service.upper(), f.severity.value, f.status.value, ctrl_name],
                [1.3, 0.7, 0.7, 0.6, 3.2], row_idx=i,
                colors=[None, None, SEV_COLORS.get(f.severity.value), status_clr, None])

        self._page_break(doc)

    def _build_references(self, doc: Document):
        t = self.t
        self._h1(doc, t["s6_title"])

        for ref in t["refs"]:
            self._h2(doc, ref["name"])
            self._body(doc, ref["desc"])
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            run = p.add_run("URL: ")
            run.font.name = 'Calibri'; run.font.size = Pt(11); run.font.bold = True
            run_url = p.add_run(ref["url"])
            run_url.font.name  = 'Calibri'
            run_url.font.size  = Pt(11)
            run_url.font.color.rgb = RGBColor(0x00, 0x56, 0xA3)
            doc.add_paragraph()